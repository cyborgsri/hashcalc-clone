import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import hashlib
import hmac
import os

try:
    from Crypto.Hash import MD4, RIPEMD
    pycrypto_available = True
except ImportError:
    pycrypto_available = False

ALGOS = [
    'md5', 'md4', 'sha1', 'sha256', 'sha384', 'sha512',
    'ripemd160'
]

def compute_hash(data_bytes, algo):
    if not data_bytes:
        return "(no data)"
    if algo == 'md4':
        if pycrypto_available:
            h = MD4.new()
            h.update(data_bytes)
            return h.hexdigest()
        else:
            return '(pip install pycryptodome)'
    elif algo == 'ripemd160':
        if pycrypto_available:
            h = RIPEMD.new()
            h.update(data_bytes)
            return h.hexdigest()
        else:
            return '(pip install pycryptodome)'
    elif algo in hashlib.algorithms_available:
        h = hashlib.new(algo)
        h.update(data_bytes)
        return h.hexdigest()
    else:
        return f'({algo.upper()} not available)'

def compute_hmac(data_bytes, key, algo):
    if not key:
        return "(empty key)"
    try:
        if algo == 'md4':
            return "(HMAC not supported for MD4)"
        if algo == 'ripemd160':
            return "(HMAC not supported for RIPEMD160)"
        h = hmac.new(key, data_bytes, algo)
        return h.hexdigest()
    except Exception as e:
        return f"(error: {e})"

class HashCalcGUI(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("HashCalc Classic Clone")
        self.geometry("900x480")
        self.minsize(700, 350)
        self.resizable(True, True)
        self.configure(bg="#f0f0f0")

        # --- Data Input Frame ---
        data_frame = ttk.LabelFrame(self, text="Data Format:")
        data_frame.pack(fill="x", pady=6, padx=16)
        self.data_var = tk.StringVar(value='File')
        for text in ['File', 'Text', 'Hex']:
            ttk.Radiobutton(data_frame, text=text, variable=self.data_var, value=text, command=self.switch_input).pack(side="left", padx=5)
        self.data_entry = ttk.Entry(data_frame)
        self.data_entry.pack(side="left", padx=8, expand=True, fill="x")
        self.browse_btn = ttk.Button(data_frame, text="Browse", command=self.browse_file)
        self.browse_btn.pack(side="left", padx=8)
        self.data_entry.bind("<Return>", self._on_enter)
        self.bind("<Return>", self._on_enter)

        # --- HMAC Section ---
        hmac_frame = ttk.LabelFrame(self, text="HMAC")
        hmac_frame.pack(fill="x", padx=16, pady=2)
        self.hmac_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(hmac_frame, text="Enable HMAC", variable=self.hmac_var).pack(side="left", padx=10)
        ttk.Label(hmac_frame, text="Key Format:").pack(side="left", padx=6)
        self.key_format = tk.StringVar(value="Text")
        ttk.Combobox(hmac_frame, textvariable=self.key_format, values=["Text", "Hex"], width=9, state="readonly").pack(side="left")
        ttk.Label(hmac_frame, text="Key:").pack(side="left", padx=6)
        self.key_entry = ttk.Entry(hmac_frame, width=24)
        self.key_entry.pack(side="left", padx=6)

        # --- Algorithms Frame ---
        hash_frame = ttk.LabelFrame(self, text="Hash Algorithms")
        hash_frame.pack(fill="both", padx=16, pady=10, expand=True)
        self.inner_frame = ttk.Frame(hash_frame)
        self.inner_frame.pack(fill="both", expand=True, padx=4, pady=4)
        self.inner_frame.grid_columnconfigure(0, weight=0)
        self.inner_frame.grid_columnconfigure(1, weight=1)

        self.algo_vars = {}
        self.hash_displays = {}
        for idx, algo in enumerate(ALGOS):
            var = tk.BooleanVar(value=True)
            chk = ttk.Checkbutton(self.inner_frame, text=algo.upper(), variable=var)
            chk.grid(row=idx, column=0, sticky="nw", padx=(10,3), pady=8)
            t = ttk.Entry(self.inner_frame, state="readonly", width=85)
            t.grid(row=idx, column=1, sticky="ew", padx=(0,12), pady=8)
            self.hash_displays[algo] = t
            self.algo_vars[algo] = var
        self.inner_frame.grid_columnconfigure(1, weight=1)

        # --- Button Frame ---
        btn_frame = ttk.Frame(self)
        btn_frame.pack(fill="x", pady=16)
        self.calc_btn = ttk.Button(btn_frame, text="Calculate", command=self.calculate)
        self.calc_btn.pack(side="left", padx=14)
        self.save_btn = ttk.Button(btn_frame, text="Save", command=self.save_result, state="disabled")
        self.save_btn.pack(side="left", padx=6)
        self.reset_btn = ttk.Button(btn_frame, text="Reset", command=self.reset_all)
        self.reset_btn.pack(side="left", padx=6)
        ttk.Button(btn_frame, text="Close", command=self.quit).pack(side="right", padx=10)
        ttk.Button(btn_frame, text="Help", command=self.show_help).pack(side="right", padx=10)

        self.switch_input()
        self.hash_results = {}
        self.bind("<Configure>", self.on_resize)

    def switch_input(self):
        mode = self.data_var.get()
        if mode == 'File':
            self.data_entry.config(state='normal')
            self.browse_btn.config(state='normal')
        else:
            self.data_entry.config(state='normal')
            self.browse_btn.config(state='disabled')
            self.data_entry.delete(0, tk.END)

    def reset_all(self):
        self.data_entry.delete(0, tk.END)
        self.key_entry.delete(0, tk.END)
        for algo in ALGOS:
            t = self.hash_displays[algo]
            t.config(state="normal")
            t.delete(0, tk.END)
            t.config(state="readonly")
        self.save_btn.config(state="disabled")

    def browse_file(self):
        path = filedialog.askopenfilename()
        if path:
            self.data_entry.delete(0, tk.END)
            self.data_entry.insert(0, path)

    def get_hmac_key(self):
        if not self.hmac_var.get():
            return None
        key_str = self.key_entry.get()
        if self.key_format.get() == "Hex":
            try:
                return bytes.fromhex(key_str)
            except Exception:
                return b""
        else:
            return key_str.encode()

    def calculate(self):
        mode = self.data_var.get()
        input_val = self.data_entry.get()
        try:
            if mode == 'File':
                if not os.path.isfile(input_val):
                    raise ValueError("Please provide a valid file path.")
                with open(input_val, 'rb') as f:
                    data_bytes = f.read()
            elif mode == 'Hex':
                data_bytes = bytes.fromhex(input_val)
            else:
                data_bytes = input_val.encode()
        except Exception as e:
            messagebox.showerror("Input Error", f"Unable to read input: {e}")
            return

        self.last_input_val = input_val
        self.last_input_type = mode

        use_hmac = self.hmac_var.get()
        hmac_key = self.get_hmac_key() if use_hmac else None
        self.hash_results.clear()
        self.save_btn.config(state="normal")
        for algo, var in self.algo_vars.items():
            display = self.hash_displays[algo]
            if var.get():
                if use_hmac and (hmac_key is not None and hmac_key != b""):
                    value = compute_hmac(data_bytes, hmac_key, algo)
                elif use_hmac and (hmac_key is None or hmac_key == b""):
                    value = "Invalid/Empty Key"
                else:
                    value = compute_hash(data_bytes, algo)
            else:
                value = ""
            display.config(state="normal")
            display.delete(0, tk.END)
            display.insert(0, value)
            display.config(state="readonly")
            self.hash_results[algo] = value if var.get() else ""

    def show_help(self):
        messagebox.showinfo(
            "Help",
            "Classic HashCalc - Now actual input and type stored in saved files!"
        )

    def _on_enter(self, event):
        self.calculate()

    def on_resize(self, event):
        for algo, t in self.hash_displays.items():
            t.config(width=max(35, int(self.winfo_width() // 15)))

    def save_result(self):
        filetypes = [
            ("Text File", "*.txt"),
            ("Word Doc", "*.docx"),
            ("PDF File", "*.pdf")
        ]
        save_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=filetypes
        )
        if not save_path:
            return
        ext = os.path.splitext(save_path)[1].lower()
        try:
            display_lines = []
            # Add input info
            display_lines.append(f"Input Type: {self.last_input_type}")
            display_lines.append(f"Input Value: {self.last_input_val}")
            display_lines.append("-" * 62)
            for algo in ALGOS:
                if self.algo_vars[algo].get():
                    val = self.hash_results.get(algo, "")
                    display_lines.append(f"{algo.upper()}: {val}")
            results_str = "\n".join(display_lines)
            if ext == ".txt":
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(results_str)
            elif ext == ".docx":
                from docx import Document
                doc = Document()
                doc.add_heading("HashCalc Results", 0)
                doc.add_paragraph(f"Input Type: {self.last_input_type}")
                doc.add_paragraph(f"Input Value: {self.last_input_val}")
                doc.add_paragraph("-" * 62)
                for algo in ALGOS:
                    if self.algo_vars[algo].get():
                        val = self.hash_results.get(algo, "")
                        doc.add_paragraph(f"{algo.upper()}: {val}")
                doc.save(save_path)
            elif ext == ".pdf":
                from reportlab.lib.pagesizes import letter
                from reportlab.pdfgen import canvas
                c = canvas.Canvas(save_path, pagesize=letter)
                width, height = letter
                c.setFont("Courier", 12)
                y = height - 40
                c.drawString(40, y, "HashCalc Results")
                y -= 25
                c.drawString(40, y, f"Input Type: {self.last_input_type}")
                y -= 20
                c.drawString(40, y, f"Input Value: {self.last_input_val}")
                y -= 20
                c.drawString(40, y, "-" * 48)
                y -= 20
                for algo in ALGOS:
                    if self.algo_vars[algo].get():
                        val = self.hash_results.get(algo, "")
                        c.drawString(40, y, f"{algo.upper()}: {val}")
                        y -= 18
                        if y < 50:
                            c.showPage()
                            y = height - 40
                            c.setFont("Courier", 12)
                c.save()
            else:
                messagebox.showerror("Save Error", "Unsupported file type!")
                return
            messagebox.showinfo("Save", "Results saved successfully!")
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save: {e}")

if __name__ == "__main__":
    app = HashCalcGUI()
    app.mainloop()
